from docx import Document
from openpyxl import load_workbook
import sys
import os

# -----------------------------
# DOCX FIND
# -----------------------------
def find_text_in_docx(input_path, target_text):
    doc = Document(input_path)
    count = 0

    # Paragraph
    for p in doc.paragraphs:
        count += p.text.count(target_text)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    count += paragraph.text.count(target_text)

    return count


# -----------------------------
# DOCX REPLACE
# -----------------------------
def replace_text_in_docx_safe(input_path, old_text, new_text, output_path=None):
    doc = Document(input_path)
    total_replaced = 0

    def replace_in_paragraph(paragraph):
        nonlocal total_replaced

        full_text = "".join(run.text for run in paragraph.runs)

        if old_text not in full_text:
            return

        occurrences = full_text.count(old_text)
        total_replaced += occurrences

        # Replace full paragraph text
        full_text = full_text.replace(old_text, new_text)

        # Reset runs
        paragraph.clear()

        # Insert single run
        paragraph.add_run(full_text)

    # Paragraphs
    for p in doc.paragraphs:
        replace_in_paragraph(p)

    # Tables
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    replace_in_paragraph(p)

    if output_path is None:
        output_path = input_path

    doc.save(output_path)
    return total_replaced



# -----------------------------
# EXCEL FIND
# -----------------------------
def find_text_in_excel(path, target):
    wb = load_workbook(path)
    count = 0

    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if isinstance(cell.value, str):
                    count += cell.value.count(target)

    return count


# -----------------------------
# EXCEL REPLACE
# -----------------------------
def replace_text_in_excel(path, old, new, output_path=None):
    wb = load_workbook(path)
    total = 0

    for sheet in wb.worksheets:
        for row in sheet.iter_rows():
            for cell in row:
                if isinstance(cell.value, str) and old in cell.value:
                    occurrences = cell.value.count(old)
                    cell.value = cell.value.replace(old, new)
                    total += occurrences

    if output_path is None:
        output_path = path

    wb.save(output_path)
    return total


# -----------------------------
# FILE ROUTING (auto-detect)
# -----------------------------
def find_auto(path, target):
    ext = os.path.splitext(path)[1].lower()

    if ext == ".docx":
        return find_text_in_docx(path, target)
    elif ext == ".xlsx":
        return find_text_in_excel(path, target)
    else:
        raise Exception(f"Unsupported file type: {ext}")


def replace_auto(path, old, new):
    ext = os.path.splitext(path)[1].lower()

    if ext == ".docx":
        return replace_text_in_docx_safe(path, old, new)
    elif ext == ".xlsx":
        return replace_text_in_excel(path, old, new)
    else:
        raise Exception(f"Unsupported file type: {ext}")


# -----------------------------
# MAIN ENTRY (UNTUK UNITY)
# -----------------------------
if __name__ == "__main__":
    mode = sys.argv[1]
    input_path = sys.argv[2]

    if mode == "find":
        target = sys.argv[3]
        result = find_auto(input_path, target)
        print(result)

    elif mode == "replace":
        old_text = sys.argv[3]
        new_text = sys.argv[4]
        replaced = replace_auto(input_path, old_text, new_text)
        print(f"REPLACED_COUNT={replaced}")

    else:
        print(f"Unknown mode: {mode}")
